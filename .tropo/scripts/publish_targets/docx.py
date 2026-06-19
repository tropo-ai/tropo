"""
publish_targets/docx.py — Docx target implementation for publish.pipeline class.

stage():  Convert extracted markdown content to a .docx file using python-docx.
          Stops at stage — no publish() (docx is an output artifact, not a deployment).

Requires: python-docx (pip install python-docx)

Output: external-work/docx/<pipeline-slug>/<filename>.docx
        or pipeline_def['output_path'] if specified.

v1.49.0 S4 — docx target greenfield implementation.

─── AGENT INSTRUCTIONS — TEMPLATE SELECTION ───────────────────────────────────

When exporting a document to Word (.docx), agents MUST select an appropriate
template. The canonical template library lives at:

    argo-os/03-design/output-templates/ms-word-templates/

Discover available templates:
    list_templates()  →  returns [{slug, label, path}, ...]

Standard templates:
    mos-word-template-internal.dotx  — Internal documents (team, partner briefings,
                                        strategy docs, working drafts not for
                                        external distribution)
    mos-word-template-external.dotx  — External documents (customer-facing,
                                        public-release, press, partner handoffs
                                        that leave MindBridge)

Protocol before calling stage():
    1. Call list_templates() to get available options.
    2. Present the options to the user with their slugs and intended audience.
    3. Ask: "Which template would you like — internal or external?"
    4. Set pipeline_def['template'] to the selected path.
    5. Set pipeline_def['title'] to the document title (drives the header field).
    6. Call stage().

The template controls: MindBridge logo, header layout, font, heading colours,
and spacing. The script sets core_properties.title so the TITLE field in the
header auto-populates on open.

────────────────────────────────────────────────────────────────────────────────
"""

from __future__ import annotations

import os
import re
import sys
import tempfile
import zipfile

_TARGETS_DIR = os.path.dirname(os.path.abspath(__file__))
_SCRIPTS_DIR = os.path.dirname(_TARGETS_DIR)
if _SCRIPTS_DIR not in sys.path:
    sys.path.insert(0, _SCRIPTS_DIR)

import time

from publish_types import StageResult, PublishTargetError
from publish_targets._shared import ensure_dated_slot

VAULT_ROOT = os.path.dirname(os.path.dirname(_SCRIPTS_DIR))

# ─── Helpers ─────────────────────────────────────────────────────────────────

def _resolve_output_path(pipeline_def: dict, publish_run_ts: str) -> str:
    """Return output .docx path using Phase 3a content-grouped convention.

    Phase 3a convention (Mike-G60 2026-05-24):
      02-outbox/<sub_system>/<sub_sub_system>/<slug>/<slug>-<timestamp>.docx

    Wrapper declares sub_system: + sub_sub_system: as explicit frontmatter fields.
    Falls back to external-work/docx/ for working-copy-export use_case (tropo-export.py
    path — distinct lineage, preserved per locked arch-spec 5a89297a).

    versions/ archive: ensure_dated_slot() moves prior <slug>-*.docx to <slug>/versions/
    on each run (mirrors web target dated-source packaging primitive).
    """
    # Working-copy-export stays on legacy path
    if pipeline_def.get('use_case') == 'working-copy-export':
        slug = re.sub(r'[^a-z0-9]+', '-', (pipeline_def.get('title', 'output') or 'output').lower()).strip('-')
        return os.path.join(VAULT_ROOT, 'external-work', 'docx', slug, f'{slug}.docx')

    # Explicit override
    output_path = pipeline_def.get('output_path', '') or pipeline_def.get('expected_output', '')
    if output_path:
        return os.path.expanduser(output_path) if output_path.startswith('~') else output_path

    # Phase 3a: content-grouped at the piece level
    slug = re.sub(r'[^a-z0-9]+', '-', (pipeline_def.get('title', 'output') or 'output').lower()).strip('-')
    sub_system = pipeline_def.get('sub_system', 'web')
    sub_sub_system = pipeline_def.get('sub_sub_system', 'agentic-builders')
    folder = os.path.join(VAULT_ROOT, '02-outbox', sub_system, sub_sub_system, slug)
    os.makedirs(folder, exist_ok=True)
    return ensure_dated_slot(folder, slug, publish_run_ts, 'docx')


def list_templates(studio_root: str | None = None) -> list[dict]:
    """Return available Word templates from the canonical template library.

    Each entry: {slug, label, audience, path}
    Call this before stage() to discover and present template choices to the user.
    """
    root = studio_root or VAULT_ROOT
    template_dir = os.path.join(root, '03-design', 'output-templates', 'ms-word-templates')
    if not os.path.isdir(template_dir):
        return []

    LABELS = {
        'mos-word-template-internal': ('Internal', 'Team, partner briefings, strategy docs, working drafts — not for external distribution'),
        'mos-word-template-external': ('External', 'Customer-facing, public-release, press, partner handoffs leaving MindBridge'),
    }

    templates = []
    for fname in sorted(os.listdir(template_dir)):
        if not fname.lower().endswith(('.dotx', '.docx')):
            continue
        slug = fname.replace('.dotx', '').replace('.docx', '')
        label, audience = LABELS.get(slug, (slug, ''))
        templates.append({
            'slug': slug,
            'label': label,
            'audience': audience,
            'path': os.path.join(template_dir, fname),
        })
    return templates


def _load_document(template_path: str | None = None):
    """Return a python-docx Document, optionally seeded from a .docx or .dotx template.

    .dotx files require a content-type patch (dotx declares template MIME type;
    python-docx only accepts document MIME type). We patch in a temp file and delete it.
    """
    from docx import Document

    if not template_path:
        return Document()

    if not os.path.exists(template_path):
        raise PublishTargetError(f'Template not found: {template_path}')

    if template_path.lower().endswith('.dotx'):
        tmp = tempfile.mktemp(suffix='.docx')
        try:
            with zipfile.ZipFile(template_path, 'r') as zin, \
                 zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == '[Content_Types].xml':
                        data = data.replace(
                            b'wordprocessingml.template.main+xml',
                            b'wordprocessingml.document.main+xml',
                        )
                    zout.writestr(item, data)
            return Document(tmp)
        finally:
            if os.path.exists(tmp):
                os.unlink(tmp)

    return Document(template_path)


def _strip_frontmatter(content: str) -> str:
    """Remove YAML frontmatter from markdown content."""
    return re.sub(r'^---\n.*?\n---\n*', '', content, count=1, flags=re.DOTALL)


def _parse_heading_level(line: str) -> tuple[int, str]:
    """Return (heading_level, text) for a markdown heading line, or (0, line)."""
    m = re.match(r'^(#{1,6})\s+(.*)', line)
    if m:
        return len(m.group(1)), m.group(2).strip()
    return 0, line


def _markdown_to_docx(content: str, title: str, doc):
    """
    Convert markdown content to docx paragraphs.
    Handles: headings (H1-H3), paragraphs, bold (**text**), bullet lists (- item).
    Sufficient for vault article content; complex tables and code blocks are flattened.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = content.splitlines()
    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # Code blocks — flatten to monospace paragraph
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.name = 'Courier New'
                p.runs[0].font.size = Pt(9)
            i += 1
            continue

        # Headings
        level, text = _parse_heading_level(line)
        if level == 1:
            doc.add_heading(text, level=1)
            i += 1
            continue
        if level == 2:
            doc.add_heading(text, level=2)
            i += 1
            continue
        if level == 3:
            doc.add_heading(text, level=3)
            i += 1
            continue

        # Bullet lists — try List Bullet, fall back to List Paragraph, then Normal with prefix
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            for style in ('List Bullet', 'List Paragraph', 'Normal'):
                try:
                    p = doc.add_paragraph(text, style=style)
                    if style == 'Normal':
                        p.runs[0].text = f'• {text}' if p.runs else text
                    break
                except KeyError:
                    continue
            i += 1
            continue

        # Numbered lists — try List Number, fall back gracefully
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            for style in ('List Number', 'List Paragraph', 'Normal'):
                try:
                    doc.add_paragraph(text, style=style)
                    break
                except KeyError:
                    continue
            i += 1
            continue

        # Horizontal rules and table rows — skip
        if line.strip() in ('---', '***', '___') or line.strip().startswith('|'):
            i += 1
            continue

        # Empty lines → paragraph break (already handled by add_paragraph)
        if not line.strip():
            i += 1
            continue

        # Regular paragraph — collect consecutive non-empty lines
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].strip().startswith('```'):
            para_lines.append(lines[i].strip())
            i += 1
        if para_lines:
            para_text = ' '.join(para_lines)
            # Strip markdown links: [text](url) → text
            para_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', para_text)
            # Strip bold/italic markers
            para_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', para_text)
            para_text = re.sub(r'\*([^*]+)\*', r'\1', para_text)
            para_text = re.sub(r'`([^`]+)`', r'\1', para_text)
            doc.add_paragraph(para_text)
        continue

    return doc


# ─── stage() ─────────────────────────────────────────────────────────────────

def stage(extracted_content: dict, pipeline_def: dict) -> StageResult:
    """
    Convert extracted markdown content to a single .docx file.

    Multiple files are concatenated in the order they appear in extracted_content,
    each preceded by its filename as a section heading (unless single-file).
    Directory entries (None values) are ignored for docx output.

    Returns StageResult with output_paths = [path_to_docx].
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError:
        raise PublishTargetError(
            'python-docx not installed. Run: pip install python-docx',
        )

    dry_run: bool = pipeline_def.get('_dry_run', False)
    verbose: bool = pipeline_def.get('_verbose', False)
    publish_run_ts = pipeline_def.get('_publish_run_ts') or time.strftime('%Y-%m-%dT%H%M%S', time.gmtime())

    output_path = _resolve_output_path(pipeline_def, publish_run_ts)
    title = pipeline_def.get('title', 'Document')

    # Filter to file entries only (skip directory sentinels)
    file_entries = {k: v for k, v in extracted_content.items() if v is not None}

    if not file_entries:
        return StageResult(
            success=False,
            errors=['No file content to convert — extracted_content contained only directory entries'],
        )

    if dry_run:
        print(f'  [docx] DRY RUN — would write {len(file_entries)} file(s) to: {output_path}')
        return StageResult(
            success=True,
            output_paths=[output_path],
            extracted_count=len(file_entries),
        )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    template_path = pipeline_def.get('template')
    doc = _load_document(template_path)
    if template_path and verbose:
        print(f'  [docx] Using template: {template_path}')

    # Set document core properties so DOCPROPERTY Title field in the header resolves
    doc.core_properties.title = title

    # Tell Word to refresh all fields on open (fixes stale DOCPROPERTY display)
    from lxml import etree
    settings = doc.settings.element
    update_fields = etree.SubElement(
        settings,
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields',
    )
    update_fields.set(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '1'
    )

    # Remove any empty stub paragraphs the template body carries (avoids leading indent/bar)
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    body = doc.element.body
    for p in list(body):
        tag = p.tag.split('}')[-1] if '}' in p.tag else p.tag
        if tag == 'p':
            has_text = bool(p.findall(f'.//{{{W}}}t'))
            if not has_text:
                body.remove(p)

    # Document title
    title_heading = doc.add_heading(title, level=0)

    multi_file = len(file_entries) > 1

    for filename, content in file_entries.items():
        if multi_file:
            section_title = re.sub(r'\.md$', '', os.path.basename(filename))
            doc.add_heading(section_title, level=1)
        clean = _strip_frontmatter(content)
        _markdown_to_docx(clean, title=filename, doc=doc)
        if verbose:
            print(f'  [docx] Added: {filename}')

    doc.save(output_path)
    print(f'  [docx] Saved: {output_path}')

    return StageResult(
        success=True,
        output_paths=[output_path],
        extracted_count=len(file_entries),
        metadata={'output_path': output_path},
    )

# No publish() — docx target stops at stage.
